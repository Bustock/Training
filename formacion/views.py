import pandas as pd
#import xlwings as xw
#from openpyxl import Workbook, load_workbook
#from openpyxl.styles import Alignment, Border, Side, PatternFill, Font
import json, base64, os, re
from .models import *
from .forms import *
from copy import copy
from django.utils import timezone
from django.db import transaction
from django.db.models import Q
from datetime import timedelta
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.contrib.messages import get_messages
from django.contrib.auth import logout, login
from django.contrib.auth.views import PasswordChangeView
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from django.http import FileResponse, Http404, HttpResponse, HttpResponseRedirect, JsonResponse
from django.urls import reverse, reverse_lazy
from django.views.decorators.csrf import csrf_exempt
from urllib.parse import urlencode
from docx import Document
from docx.shared import Inches
from PIL import Image
from datetime import datetime
from io import BytesIO
from fpdf import FPDF
from PyPDF2 import PdfMerger
import shutil


# BASE_URL = r"C:\sonova\formaciones"
#BASE_URL = getattr(settings, 'TRAINING_SHARED_ROOT', r"D:\Training")
BASE_URL = getattr(settings, 'TRAINING_SHARED_ROOT', r"\\es01sw31\APP Training Tool")


def shared_path(*parts):
    return os.path.join(BASE_URL, *parts)


def shared_media_path(*parts):
    return shared_path('media', *parts)


def shared_plantillas_path(*parts):
    return shared_media_path('plantillas', *parts)


def json_error_detallado(error, fase, puesto_seleccionado=None, status=500):
    ruta = None
    if puesto_seleccionado:
        ruta = shared_plantillas_path(fase, f"{puesto_seleccionado}.txt")

    return JsonResponse({
        'ok': False,
        'fase': fase,
        'puesto': puesto_seleccionado,
        'ruta_esperada': ruta,
        'detalle_error': str(error),
    }, status=status)


opi_a_mod = []
puesto_info = []

puestos_dict = {
    'OPERARIO': 'OPERARIO',
    'CONFIRM': 'CONFIRMACIÓN',
    'OPEN_MAIL': 'OPEN MAIL',
    'ORDER_ENTRY_SERVICE': 'ORDER ENTRY SERVICE',
    'ORDER_ENTRY_NEW': 'ORDER ENTRY NEW',
    'DIAG_ITE': 'DIAGNÓSTICO ITE',
    'DIAG_ITE_SAF': 'DIAGNÓSTICO ITE (SAF)',
    'DIAG_SDS': 'DIAGNÓSTICO SDS',
    'DIAG_SDS_SAF': 'DIAGNÓSTICO SDS (SAF)',
    'DIAG_BTE': 'DIAGNÓSTICO BTE',
    'DIAG_FM_DWA': 'DIAGNÓSTICO FM/DWA',
    'DIAG_TITANIUM': 'DIAGNÓSTICO TITANIUM',
    'REPAIR_FM_DWA': 'REPAIR FM/DWA',
    'REPAIR_CARGADORES': 'REPAIR CARGADORES',
    'TRESA_BLOQUE_1': 'TRESA BLOQUE 1',
    'TRESA_BLOQUE_2': 'TRESA BLOQUE 2',
    'TRESA_BLOQUE_3': 'TRESA BLOQUE 3',
    'REPAIR_BTE': 'REPAIR BTE',
    'RSM': 'RSM',
    'MINI_KITTING': 'MINI-KITTING',
    'KITTING_CUSTOM': 'KITTING CUSTOM',
    'KITTING_BTE': 'KITTING BTE',
    'DCC': 'DCC',
    'DLP': 'DLP',
    'KIT_PREP_SDS': 'KIT PREPARATION SDS',
    'CLOSING_SDS': 'CLOSING SDS',
    'KIT_PREP_ITE': 'KIT PREPARATION ITE',
    'CLOSING_ITE': 'CLOSING ITE',
    'CUT_TRIM': 'CUT & TRIM',
    'FL': 'FINISHING & LACQUERING',
    'TITANIUM': 'TITANIUM',
    'VISUAL_ITE': 'VISUAL ITE',
    'VISUAL_SDS': 'VISUAL SDS',
    'VISUAL_BTE': 'VISUAL BTE',
    'VISUAL_FM_DWA': 'VISUAL FM/DWA',
    'PRO_GO_ITE': 'PRO&GO ITE',
    'PRO_GO_SDS': 'PRO&GO SDS',
    'PRO_GO_BTE': 'PRO&GO BTE',
    'PRO_GO_FM_DWA': 'PRO&GO FM/DWA',
    'PACKING_NEW': 'PACKING NEW',
    'BULK': 'BULK',
    'PACKING_SERVICE': 'PACKING SERVICE',
    'RPM_DESMONTAJE_LIMPIEZA': 'RPM: DESMONTAJE Y LIMPIEZA',
    'RPM_DECONTAMINACION': 'RPM: DECONTAMINACIÓN',
    'RPM_IRS1_IDENTIFICACION': 'RPM-IRS1: IDENTIFICACIÓN E INICIALIZACIÓN',
    'RPM_LPS2_BLUETOOTH': 'RPM-LPS2: BLUETOOTH',
    'RPM_VISUAL_SNW2': 'RPM-VISUAL Y SNW2',
    'ACS2': 'ACS2',
    'SORTING': 'SORTING',
    'REFURBISHING': 'REFURBISHING',
    'REPROCESSING_CARGADORES': 'REPROCESSING CARGADORES',
    'PACKING_RPM_REFURBISHING': 'PACKING RPM/REFURBISHING'
}

def registro(request):
    if request.method == "POST":
        form = RegistroForm(request.POST)
        if form.is_valid():
            user = form.save()
            PasswordChangeStatus.objects.update_or_create(
                user=user,
                defaults={"last_password_change": timezone.now()},
            )
            grupo_tecnicos, creado = Group.objects.get_or_create(name="tecnicos")
            user.groups.add(grupo_tecnicos)
            login(request, user, backend='django.contrib.auth.backends.ModelBackend')  # inicia sesión automáticamente después de registrarse
            return redirect("login")  # cámbialo a la vista principal de tu app
    else:
        form = RegistroForm()
    return render(request, "registro.html", {"form": form})

def logout_message(request):
    # Recuperar mensaje de firma de la sesión si lo hemos guardado
    mensaje_firma = request.session.pop('mensaje_firma', None)
    logout(request)
    if mensaje_firma:
        messages.info(request, mensaje_firma)
    else:
        messages.info(request, "Sesión cerrada correctamente.")
    
    return redirect('login')

def timeout(request):
    messages.info(request, "Sesión expirada. Por favor, inicia sesión de nuevo.")
    return redirect('login')

def group_required(group_name):
    def in_group(u):
        return u.is_authenticated and u.groups.filter(name=group_name).exists()
    return user_passes_test(in_group)

def groups_required(*group_names):
    def in_groups(u):
        return u.is_authenticated and bool(u.groups.filter(name__in=group_names))
    return user_passes_test(in_groups)


class CambioPasswordView(PasswordChangeView):
    template_name = "registration/password_change_form.html"
    success_url = reverse_lazy("password_change_done")

    def form_valid(self, form):
        response = super().form_valid(form)
        PasswordChangeStatus.objects.update_or_create(
            user=self.request.user,
            defaults={"last_password_change": timezone.now()},
        )
        return response

@groups_required('admin', 'formacion', 'supervisores')
@login_required
def inicio(request):
    if request.user.groups.filter(name='supervisores').exists():
        return redirect('listar_opis')
    else:
        notificaciones = Notificacion.objects.filter(
            models.Q(grupo__in=request.user.groups.all()) |
            models.Q(usuario=request.user)
        ).exclude(leido_por=request.user)
        puede_plantillas = request.user.groups.filter(name__in=['admin', 'formacion']).exists()
        return render(request, 'inicio.html', {
            'notificaciones': notificaciones,
            'puede_plantillas': puede_plantillas,
            'mensaje': messages.get_messages(request)})

@login_required    
def notificaciones(request):
    notificaciones = Notificacion.objects.filter(
        models.Q(grupo__in=request.user.groups.all()) |
        models.Q(usuario=request.user)
    ).exclude(leido_por=request.user)
    
    return render(request, 'notificaciones.html', {
        'notificaciones': notificaciones})

@login_required
def notificacion_leida(request, notificacion_id):
    notificacion = get_object_or_404(Notificacion, id=notificacion_id)
    notificacion.leido_por.add(request.user)
    return redirect('notificaciones')  # Ajusta al nombre de tu url de notificaciones

@groups_required('admin', 'formacion')
@login_required
def actualizar_matriz(request):
    # Usar rutas relativas al proyecto
    base_dir = settings.BASE_DIR
    excel_file = os.path.join(base_dir, 'TPL-708 [4] RCSE Matriz de Polivalencia Operarios producción.xlsx')

    # Mapeo de columnas entre el Excel y los campos de la base de datos
    column_mapping = {
        'Operario': 'OPERARIO',
        'CONFIRM': 'CONFIRM',
        'OPEN MAIL': 'OPEN_MAIL',
        'ORDER ENTRY SERVICE': 'ORDER_ENTRY_SERVICE',
        'ORDER ENTRY NEW': 'ORDER_ENTRY_NEW',
        'DIAG. ITE': 'DIAG_ITE',
        'DIAG. ITE SAF': 'DIAG_ITE_SAF',
        'DIAG. SDS': 'DIAG_SDS',
        'DIAG. SDS SAF': 'DIAG_SDS_SAF',
        'DIAG. BTE': 'DIAG_BTE',
        'DIAG.FM/DWA': 'DIAG_FM_DWA',
        'DIAG. TITANIUM': 'DIAG_TITANIUM',
        'REPAIR FM/DWA': 'REPAIR_FM_DWA',
        'REPAIR CARGADORES': 'REPAIR_CARGADORES',
        'TRESA Bloque 1': 'TRESA_BLOQUE_1',
        'TRESA Bloque 2': 'TRESA_BLOQUE_2',
        'TRESA Bloque 3': 'TRESA_BLOQUE_3',
        'REPAIR BTE': 'REPAIR_BTE',
        'RSM': 'RSM',
        'MINI-KITTING': 'MINI_KITTING',
        'KITTING CUSTOM': 'KITTING_CUSTOM',
        'KITTING BTE': 'KITTING_BTE',
        'DCC': 'DCC',
        'DLP': 'DLP',
        'KIT PREP. SDS': 'KIT_PREP_SDS',
        'CLOSING SDS': 'CLOSING_SDS',
        'KIT PREP. ITE': 'KIT_PREP_ITE',
        'CLOSING ITE': 'CLOSING_ITE',
        'CUT & TRIM': 'CUT_TRIM',
        'F&L': 'FL',
        'TITANIUM': 'TITANIUM',
        'VISUAL ITE': 'VISUAL_ITE',
        'VISUAL SDS': 'VISUAL_SDS',
        'VISUAL BTE': 'VISUAL_BTE',
        'VISUAL FM/DWA': 'VISUAL_FM_DWA',
        'PRO& GO ITE': 'PRO_GO_ITE',
        'PRO& GO SDS': 'PRO_GO_SDS',
        'PRO& GO BTE': 'PRO_GO_BTE',
        'PRO & GO FM/DWA ': 'PRO_GO_FM_DWA',
        'PACKING NEW': 'PACKING_NEW',
        'BULK': 'BULK',
        'PACKING SERVICE': 'PACKING_SERVICE',
        'RPM-Desmontaje y limpieza': 'RPM_DESMONTAJE_LIMPIEZA',
        'RPM-Descontaminación': 'RPM_DECONTAMINACION',
        'RPM-IRS1 Identificación e Inicialización': 'RPM_IRS1_IDENTIFICACION',
        'RPM-LPS2 Bluetooth': 'RPM_LPS2_BLUETOOTH',
        'RPM-Visual y SNW2': 'RPM_VISUAL_SNW2',
        'ACS2': 'ACS2',
        'Sorting': 'SORTING',
        'Refurbishing': 'REFURBISHING',
        'Reprocessing cargadores': 'REPROCESSING_CARGADORES',
        'Packing RPM/Refurbishing': 'PACKING_RPM_REFURBISHING'
    }

    # Verificar que el archivo Excel existe
    if not os.path.exists(excel_file):
        messages.add_message(request, messages.ERROR, f'No se encuentra el archivo Excel en: {excel_file}')
        return redirect(inicio)

    try:
        # Leer el archivo Excel desde la hoja "Operarios producción", comenzando desde la fila 3
        df = pd.read_excel(excel_file, sheet_name="Operarios producción", skiprows=2)
        df.fillna(0, inplace=True)
        df.rename(columns=column_mapping, inplace=True)
    except Exception as e:
        messages.add_message(request, messages.ERROR, f'Error al leer el archivo Excel: {str(e)}')
        return redirect(inicio)

    campos_matriz = [campo for campo in column_mapping.values() if campo != 'OPERARIO']
    usa_escala_legacy = False
    for campo in campos_matriz:
        # Si la columna no existe en el Excel, usar una serie vacía para evitar errores.
        serie = pd.to_numeric(df.get(campo, pd.Series(dtype='float64')), errors='coerce')
        if (serie.dropna().astype(int) == 4).any():
            usa_escala_legacy = True
            break

    def normalizar_nivel(valor):
        if pd.isna(valor):
            return 0
        try:
            nivel = int(valor)
        except (TypeError, ValueError):
            return 0

        # Compatibilidad con escala legacy: 4=en formación, 1/2=formado.
        if usa_escala_legacy:
            if nivel == 4:
                return 1
            if nivel in [1, 2]:
                return 2
            if nivel == 3:
                return 3
            return 0

        if nivel in [0, 1, 2, 3]:
            return nivel
        if nivel == 4:
            return 1
        return 0

    operarios_validos = False
    operarios_vistos = set()
    polivalencias_a_crear = []

    for _, row in df.iterrows():
        operario = row.get('OPERARIO')
        if not operario or pd.isna(operario):
            continue

        operario = str(operario).strip()
        if operario.upper() == 'LEYENDA':
            print("Fila 'LEYENDA' detectada. Fin de operarios.")
            break
        if operario in operarios_vistos:
            print(f"Operario duplicado encontrado en el Excel: {operario}")
            continue

        operarios_vistos.add(operario)
        operarios_validos = True

        datos_polivalencia = {'OPERARIO': operario}
        for campo in campos_matriz:
            datos_polivalencia[campo] = normalizar_nivel(row.get(campo, 0))

        polivalencias_a_crear.append(polivalencia(**datos_polivalencia))

    if not operarios_validos:
        messages.add_message(request, messages.ERROR, 'No se encontraron operarios válidos en el Excel.')
        return redirect(inicio)

    campos_puesto = campos_matriz

    with transaction.atomic():
        polivalencia.objects.all().delete()
        polivalencia.objects.bulk_create(polivalencias_a_crear, batch_size=1000)

        matriz_actual = list(polivalencia.objects.all())
        mapa_matriz = {
            tecnico.OPERARIO: {campo: getattr(tecnico, campo, 0) for campo in campos_puesto}
            for tecnico in matriz_actual
        }

        existentes_completa = set(completa.objects.values_list('OPERARIO', 'PUESTO'))
        nuevas_combinaciones = []
        for tecnico in matriz_actual:
            for puesto in campos_puesto:
                if getattr(tecnico, puesto, 0) == 1 and (tecnico.OPERARIO, puesto) not in existentes_completa:
                    nuevas_combinaciones.append(
                        completa(
                            PUESTO=puesto,
                            OPERARIO=tecnico.OPERARIO,
                            TEORIA=False,
                            PRACTICA=False,
                            PRODUCTO=False,
                            firmas={}
                        )
                    )
                    existentes_completa.add((tecnico.OPERARIO, puesto))

        if nuevas_combinaciones:
            completa.objects.bulk_create(nuevas_combinaciones, batch_size=1000)

        ids_a_eliminar = []
        advertencias = []
        for registro in completa.objects.all():
            valor_actual = mapa_matriz.get(registro.OPERARIO, {}).get(registro.PUESTO)
            if valor_actual in [2, 3]:
                if registro.TEORIA and registro.PRACTICA and registro.PRODUCTO:
                    ids_a_eliminar.append(registro.id)
                else:
                    advertencias.append(
                        f"⚠️ El operario {registro.OPERARIO} no ha completado la formación en {registro.PUESTO}, pero ha cambiado su valor en la matriz de polivalencia.⚠️"
                    )

        if ids_a_eliminar:
            completa.objects.filter(id__in=ids_a_eliminar).delete()

    for advertencia in advertencias:
        messages.add_message(request, messages.INFO, advertencia)

    messages.add_message(request, messages.INFO, 'Matriz de polivalencia actualizada.')
    return redirect(inicio)


##################################################################################### DEF PARA LAS OPIS ######################################################################################

@groups_required('admin', 'formacion')
@login_required
def formacion_opis(request):
    global opi_a_mod
    global puestos_dict
    operario_form = OperarioForm(request.GET)
    operarios_info = []
    puesto_form = PuestoForm(request.GET)
    tecnicos_info = []
    puesto_info = []
    opi_form = OpiForm(request.GET)
    opis_info = []
    formados_info = []
    sin_firma_info = []
    firma_info = []
    puesto_a_buscar = []
    puesto_seleccionado = []
    puesto_traducido = []

    if opi_form.is_valid():
        opi = opi_form.cleaned_data.get('opi')

        if opi:
            opi_obj = opis.objects.filter(OPI=opi).first()
            opi_a_mod = opi_obj
            secciones = [opi_obj.SECCION1, opi_obj.SECCION2, opi_obj.SECCION3]

            formados_dict = opi_obj.formados if isinstance(opi_obj.formados, dict) else {}
            firmas_dict = opi_obj.firmas if isinstance(opi_obj.firmas, dict) else {}

            for nombre, fecha in formados_dict.items():
                formados_info.append({'nombre': nombre, 'fecha': fecha})

            for operario in polivalencia.objects.all():
                if operario.OPERARIO in formados_dict:
                    if operario.OPERARIO not in firmas_dict:
                        sin_firma_info.append(operario.OPERARIO)
                    elif operario.OPERARIO in firmas_dict:
                        firma_info.append({'nombre': operario.OPERARIO, 'firma': firmas_dict[operario.OPERARIO]})
                    continue

                valores_validos = {}
                for seccion in secciones:
                    if hasattr(operario, seccion):
                        field_value = getattr(operario, seccion)
                        if isinstance(field_value, int) and field_value not in [0, 1]:
                            valores_validos[seccion] = field_value

                if valores_validos:
                    opis_info.append({'nombre': operario.OPERARIO})

    if operario_form.is_valid():
        operario_seleccionado = operario_form.cleaned_data.get('OPERARIO')
        if operario_seleccionado:
            operario_obj = polivalencia.objects.filter(OPERARIO=operario_seleccionado).first()
            if operario_obj:
                campos_excluidos = ['id', 'OPERARIO', 'creado_por', 'creado_en', 'modificado_por', 'modificado_en']
                campos_validos = [
                    (field.name, puestos_dict.get(field.name, field.name))
                    for field in polivalencia._meta.fields
                    if field.name not in campos_excluidos and getattr(operario_obj, field.name) not in [0, 1]
                ]

                operarios_info = {
                    'nombre': operario_obj.OPERARIO,
                    'campos_validos': campos_validos
                }

    if puesto_form.is_valid():
        puesto_info.clear()
        puesto_seleccionado = puesto_form.cleaned_data.get('PUESTO')

        if puesto_seleccionado:
            puesto_traducido = puestos_dict[puesto_seleccionado] 
            # Filtrar técnicos que NO tengan 0 ni 1 en el campo correspondiente
            tecnicos_info = polivalencia.objects.exclude(
                Q(**{puesto_seleccionado: 0}) | Q(**{puesto_seleccionado: 1})
            ).values('OPERARIO', puesto_seleccionado)

            # Crear un filtro dinámico con Q objects para SECCION1 a SECCION7
            filtro = Q()
            for i in range(1, 8):
                filtro |= Q(**{f'SECCION{i}': puesto_seleccionado})

            puesto_a_buscar = opis.objects.filter(filtro)

            # Agregar las OPIs encontradas en opis_info a puesto_info
            puesto_info.extend(opi.OPI for opi in puesto_a_buscar)


    return render(request, 'form_opis.html', {
        'opi_form': opi_form,
        'opis_info': opis_info,
        'opi_a_mod': opi_a_mod,
        'formados_info': formados_info,
        'sin_firma_info': sin_firma_info,
        'firma_info': firma_info,
        'operario_form': operario_form,
        'operarios_info': operarios_info,
        'puesto_form': puesto_form,
        'tecnicos_info': tecnicos_info,
        'puesto_info': puesto_info,
        'puesto_seleccionado': puesto_seleccionado,
        'puesto_traducido': puesto_traducido,
        'mensaje': messages.get_messages(request)
    })

@groups_required('admin', 'formacion')
@login_required
def nueva_opi(request):
    puesto_form = SeccionForm(request.GET)
    tecnicos_info = []

    # Si el formulario del puesto es válido y se ha enviado
    if puesto_form.is_valid():
        puesto = puesto_form.cleaned_data.get('puesto')
        
        if puesto:
            for operario in polivalencia.objects.all():
                field_value = getattr(operario, puesto)
                if field_value != 0:
                    tecnicos_info.append({
                        'nombre': operario.OPERARIO,
                        'puesto': puesto,
                        'valor': field_value
                    })

    return render(request, 'nueva_opi.html', {
        'puesto_form': puesto_form,
        'tecnicos_info': tecnicos_info,
    })

@groups_required('admin', 'formacion')
@login_required
def guardar_opi(request):
    if request.method == 'POST':
        TIPO = request.POST['tipo'].upper()
        ID = request.POST['ID'].upper()
        VERSION = request.POST['version'].upper()
        INFO = request.POST['INFO'].upper()
        SECCION1 = request.POST['SECCION1']
        SECCION2 = request.POST['SECCION2']
        SECCION3 = request.POST['SECCION3']
        SECCION4 = request.POST['SECCION4']
        SECCION5 = request.POST['SECCION5']
        SECCION6 = request.POST['SECCION6']
        SECCION7 = request.POST['SECCION7']

        if not TIPO:
            TIPO = 'OPI'

        OPI = f"{TIPO}-{ID} [{VERSION}]"

    datos = nuevas_opis(OPI=OPI, INFO=INFO, SECCION1=SECCION1, SECCION2=SECCION2, SECCION3=SECCION3, SECCION4=SECCION4, SECCION5=SECCION5, SECCION6=SECCION6, SECCION7=SECCION7)
    datos.creado_por = request.user
    datos.save()
    
    messages.add_message(request, messages.INFO, 'OPI guardada correctamente, Esperando la validación del supervisor.')
    return redirect(formacion_opis)

@groups_required('admin', 'supervisores')
@login_required
def listar_opis(request):
    global puestos_dict

    opis_qs = nuevas_opis.objects.all().order_by('OPI')

    opis_lista = []
    for opi in opis_qs:
        ok_dict = opi.ok_supervisor or {}

        secciones_aceptadas = []
        secciones_rechazadas = []
        secciones_pendientes = []

        for seccion_field in ['SECCION1', 'SECCION2', 'SECCION3', 'SECCION4', 'SECCION5', 'SECCION6', 'SECCION7']:
            clave = getattr(opi, seccion_field)
            if clave:
                valor = puestos_dict.get(clave, clave)  # traducir si existe, o mostrar la clave
                if clave in ok_dict and ok_dict[clave] == 'ok':
                    secciones_aceptadas.append(valor)
                elif clave in ok_dict and ok_dict[clave] == 'ko':
                    secciones_rechazadas.append(valor)
                else:
                    secciones_pendientes.append(valor)

        opis_lista.append({
            'nombre': opi.OPI,
            'secciones_aceptadas': secciones_aceptadas,
            'secciones_rechazadas': secciones_rechazadas,
            'secciones_pendientes': secciones_pendientes,
        })

    return render(request, 'listar_opis.html', {
        'opis_lista': opis_lista,
    })

@login_required
@groups_required('admin', 'supervisores')
def aceptar_opi(request):
    global puestos_dict

    opi_nombre = request.GET.get('opi_nombre')
    puesto_traducido = request.GET.get('puesto')

    if not opi_nombre or not puesto_traducido:
        messages.error(request, 'Faltan datos para procesar la aceptación.')
        return redirect('listar_opis')

    clave_puesto = next((clave for clave, valor in puestos_dict.items() if valor == puesto_traducido), None)

    if not clave_puesto:
        messages.error(request, f'No se encontró la clave asociada al puesto "{puesto_traducido}".')
        return redirect('listar_opis')

    opi_obj = get_object_or_404(nuevas_opis, OPI=opi_nombre)

    ok_dict = opi_obj.ok_supervisor or {}
    ok_dict[clave_puesto] = 'ok'
    opi_obj.ok_supervisor = ok_dict
    opi_obj.modificado_por = request.user
    opi_obj.save()

    # Obtener secciones asociadas y gestionadas
    secciones_asociadas = [
        getattr(opi_obj, f'SECCION{i}') for i in range(1, 8)
        if getattr(opi_obj, f'SECCION{i}')
    ]
    todas_gestionadas = all(seccion in ok_dict for seccion in secciones_asociadas)

    if todas_gestionadas:
        # Compactar secciones aceptadas en orden
        secciones_aceptadas = [
            seccion for seccion in secciones_asociadas if ok_dict.get(seccion) == 'ok'
        ]
        secciones_rechazadas = [
            seccion for seccion in secciones_asociadas if ok_dict.get(seccion) == 'ko'
        ]
        secciones_rechazadas_traducidas = [puestos_dict.get(seccion, seccion) for seccion in secciones_rechazadas]
        if secciones_aceptadas == secciones_asociadas:
            messages.success(request, f'{opi_nombre} aceptada para puesto {puesto_traducido}.')
            Notificacion.objects.create(
            grupo=Group.objects.get(name='formacion'),  # grupo
            creado_por=request.user,
            mensaje=f'✅{opi_nombre} aceptada para puesto {puesto_traducido}✅'
            )
            Notificacion.objects.create(
                grupo=Group.objects.get(name='formacion'),  # grupo
                creado_por=request.user,
                mensaje=f'✅{opi_nombre} ACEPTADA por todos✅'
            )
        else:
            messages.success(request, f'{opi_nombre} aceptada para puesto {puesto_traducido}.')
            Notificacion.objects.create(
                grupo=Group.objects.get(name='formacion'),  # grupo
                creado_por=request.user,
                mensaje=f'✅{opi_nombre} ACEPTADA parcialmente, ❌ Rechazada en secciones: {", ".join(secciones_rechazadas_traducidas)}'
            )
        campos_seccion = {}
        for idx in range(1, 8):
            if idx <= len(secciones_aceptadas):
                campos_seccion[f'SECCION{idx}'] = secciones_aceptadas[idx - 1]
            else:
                campos_seccion[f'SECCION{idx}'] = ''

        opi_nueva = opis.objects.create(
            OPI=opi_obj.OPI,
            INFO=opi_obj.INFO,
            formados={},
            firmas={},
            creado_por=request.user,
            **campos_seccion
        )
        opi_obj.delete()
    else:
        messages.success(request, f'{opi_nombre} aceptada para puesto {puesto_traducido}.')
        Notificacion.objects.create(
            grupo=Group.objects.get(name='formacion'),  # grupo
            creado_por=request.user,
            mensaje=f'✅{opi_nombre} aceptada para puesto {puesto_traducido}✅'
        )

    return redirect('listar_opis')

@login_required
@groups_required('admin', 'supervisores')
def rechazar_opi(request):
    global puestos_dict

    opi_nombre = request.GET.get('opi_nombre')
    puesto_traducido = request.GET.get('puesto')

    if not opi_nombre or not puesto_traducido:
        messages.error(request, 'Faltan datos para procesar el rechazo.')
        return redirect('listar_opis')

    clave_puesto = next((clave for clave, valor in puestos_dict.items() if valor == puesto_traducido), None)

    if not clave_puesto:
        messages.error(request, f'No se encontró la clave asociada al puesto "{puesto_traducido}".')
        return redirect('listar_opis')

    opi_obj = get_object_or_404(nuevas_opis, OPI=opi_nombre)

    ok_dict = opi_obj.ok_supervisor or {}
    ok_dict[clave_puesto] = 'ko'
    opi_obj.ok_supervisor = ok_dict
    opi_obj.modificado_por = request.user
    opi_obj.save()

    # Obtener secciones asociadas y gestionadas
    secciones_asociadas = [
        getattr(opi_obj, f'SECCION{i}') for i in range(1, 8)
        if getattr(opi_obj, f'SECCION{i}')
    ]
    todas_gestionadas = all(seccion in ok_dict for seccion in secciones_asociadas)

    if todas_gestionadas:
        messages.success(request, f'{opi_nombre} rechazada para puesto {puesto_traducido}.')
        Notificacion.objects.create(
            grupo=Group.objects.get(name='formacion'),  # grupo
            creado_por=request.user,
            mensaje=f'❌{opi_nombre} RECHAZADA para puesto {puesto_traducido}❌'
        )
        # Compactar secciones aceptadas en orden
        secciones_aceptadas = [
            seccion for seccion in secciones_asociadas if ok_dict.get(seccion) == 'ok'
        ]
        secciones_rechazadas = [
            seccion for seccion in secciones_asociadas if ok_dict.get(seccion) == 'ko'
        ]
        if not secciones_aceptadas:
            opi_obj.delete()
            Notificacion.objects.create(
            grupo=Group.objects.get(name='formacion'),  # grupo
            creado_por=request.user,
            mensaje=f'❌{opi_nombre} RECHAZADA por TODOS, eliminada del sistema❌'
            )
            return redirect('listar_opis')
        else:
            Notificacion.objects.create(
                grupo=Group.objects.get(name='formacion'),  # grupo
                creado_por=request.user,
                mensaje=f'✅{opi_nombre} ACEPTADA parcialmente, ❌ Rechazada en secciones: {", ".join(secciones_rechazadas)}'
            )
        campos_seccion = {}
        for idx in range(1, 8):
            if idx <= len(secciones_aceptadas):
                campos_seccion[f'SECCION{idx}'] = secciones_aceptadas[idx - 1]
            else:
                campos_seccion[f'SECCION{idx}'] = ''

        opi_nueva = opis.objects.create(
            OPI=opi_obj.OPI,
            INFO=opi_obj.INFO,
            formados={},
            firmas={},
            creado_por=request.user,
            **campos_seccion
        )
        opi_obj.delete()
    else:
        messages.success(request, f'{opi_nombre} rechazada para puesto {puesto_traducido}.')
        Notificacion.objects.create(
            grupo=Group.objects.get(name='formacion'),  # grupo
            creado_por=request.user,
            mensaje=f'❌{opi_nombre} RECHAZADA para puesto {puesto_traducido}❌'
        )

    return redirect('listar_opis')

@groups_required('admin', 'formacion')
@login_required
def introducir_fecha(request):
    global opi_a_mod
    opi = opi_a_mod  # ID de la OPI que se modificará
    operario = request.GET.get('operario')  # Nombre del operario

    return render(request, 'introducir_fecha.html', {'opi': opi, 'operario': operario})

@groups_required('admin', 'formacion')
@login_required
def guardar_fecha(request):
    global opi_a_mod
    if request.method == "POST":
        opi_id = opis.objects.filter(OPI=opi_a_mod).first()
        opi_id = opi_id.id
        operario = request.POST.get("operario")
        fecha = request.POST.get("fecha")

        # Buscar la OPI a modificar
        opi = get_object_or_404(opis, id=opi_id)

        # Actualizar el diccionario en el campo JSON
        if not opi.formados:
            opi.formados = {}  # Crear un diccionario si está vacío

        opi.formados[operario] = fecha  # Agregar/modificar operario:fecha

        opi.modificado_por = request.user
        opi.save()  # Guardar cambios en la BBDD

        messages.add_message(request, messages.INFO, f'{operario} completó: {opi_a_mod}')
        return redirect(f"{reverse('formacion_opis')}?opi={opi_a_mod}")  # Redirigir a la página anterior

@groups_required('admin', 'formacion', 'tecnicos')
@login_required
def introducir_firma(request):
    global opi_a_mod
    opi = opi_a_mod  # ID de la OPI que se modificará
    operario = request.GET.get('operario')  # Nombre del operario

    return render(request, 'introducir_firma.html', {'opi': opi, 'operario': operario})

def convertir_docx_a_pdf(docx_path, pdf_path):
    import subprocess

    try:
        output_dir = os.path.dirname(pdf_path)
        libreoffice_path = getattr(settings, 'LIBREOFFICE_PATH', 'soffice')

        result = subprocess.run(
            [libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60
        )

        # LibreOffice genera el PDF con el mismo nombre base que el .docx
        generated_pdf = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
        if generated_pdf != pdf_path and os.path.exists(generated_pdf):
            os.rename(generated_pdf, pdf_path)

        if result.returncode != 0:
            print(f"Error al convertir a PDF: {result.stderr.decode()}")
    except Exception as e:
        print(f"Error al convertir a PDF: {e}")

@groups_required('admin', 'formacion', 'tecnicos')
@login_required
def subir_firma(request):
    global opi_a_mod
    if request.method == "POST":
        opi_id = opis.objects.filter(OPI=opi_a_mod).first()
        opi_id = opi_id.id
        info_formacion = opis.objects.get(id=opi_id).INFO
        operario = request.POST.get("operario")
        imagen_base64 = request.POST.get("imagen")
        dni = request.POST.get("dni").strip().upper() 
        formador = (request.POST.get("formador") or "").strip().upper()

        # Buscar la OPI a modificar
        opi = get_object_or_404(opis, id=opi_id)

        # Obtener la fecha del diccionario formados
        fecha = opi.formados.get(operario, '')
        # Convertir la cadena a objeto datetime
        fecha = datetime.strptime(fecha, "%Y-%m-%d")
        # Formatear la fecha en el formato deseado
        fecha = fecha.strftime("%d-%b-%Y")

        # Obtener la hora actual
        hora = datetime.now().strftime("%H:%M")

        # Decodificar la imagen base64 y almacenarla en memoria (BytesIO)
        try:
            format, imgstr = imagen_base64.split(';base64,')
            img_bytes = base64.b64decode(imgstr)
            img_stream = BytesIO(img_bytes)  # Mantener la imagen en memoria

            # Rellenar la plantilla .docx
            # template_path = r"C:\sonova\formaciones\media\plantillas\plantilla_opis.docx"
            template_path = shared_plantillas_path('plantilla_opis.docx')
            doc = Document(template_path)

            # Función para reemplazar texto en párrafos y manejar imágenes
            def process_paragraph(paragraph, replacements, image_stream):
                text = paragraph.text
                if '{{imagen}}' in text:
                    # Separar el texto antes y después de {{imagen}}
                    parts = text.split('{{imagen}}')
                    paragraph.clear()

                    # Restaurar el texto antes de la imagen (si hay algo)
                    if parts[0]:
                        paragraph.add_run(parts[0])

                    # Insertar la imagen desde la memoria
                    run = paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(2.5))

                    # Restaurar el texto después de la imagen (si hay algo)
                    if len(parts) > 1 and parts[1]:
                        paragraph.add_run(parts[1])
                else:
                    # Reemplazar otras etiquetas normalmente
                    for key, value in replacements.items():
                        text = text.replace(key, str(value))
                    paragraph.text = text

            # Diccionario con los valores a reemplazar
            replacements = {
                '{{fecha}}': fecha,
                '{{opi_a_mod}}': opi_a_mod,
                '{{hora}}': hora,
                '{{info_formacion}}': info_formacion,
                '{{operario}}': operario,
                '{{dni}}': dni,
                '{{formador}}': formador
            }

            # Aplicar reemplazo en párrafos
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph, replacements, img_stream)

            # Aplicar reemplazo en tablas también
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph, replacements, img_stream)

            # Guardar el documento modificado
            # output_directory = os.path.join(settings.MEDIA_ROOT, 'firmas_opis', str(opi_a_mod))
            output_directory = shared_media_path('firmas_opis', str(opi_a_mod))
            os.makedirs(output_directory, exist_ok=True)
            output_path = os.path.join(output_directory, f"{operario}_acta_formacion_{opi_a_mod}.docx")
            doc.save(output_path) # Guardar el archivo .docx
            pdf_output_path = output_path.replace(".docx", ".pdf")
            convertir_docx_a_pdf(output_path, pdf_output_path) # Convertir a PDF
            os.remove(output_path)  # Eliminar el archivo .docx

            # Actualizar el diccionario en el campo JSON
            if not opi.firmas:
                opi.firmas = {}  # Crear un diccionario si está vacío

            opi.firmas[operario] = output_path  # Agregar/modificar nombre:ruta_imagen

            opi.modificado_por = request.user
            opi.save()  # Guardar cambios en la BBDD

            messages.add_message(request, messages.INFO, f'Firma de {operario} guardada en la OPI: {opi_a_mod}. Hasta la próxima!')
        except ValueError as e:
            print(f"Error al decodificar la imagen base64: {e}")
            messages.add_message(request, messages.ERROR, 'Error al guardar la firma. Por favor, inténtalo de nuevo.')
        if request.user.groups.filter(name='tecnicos').exists():
            return redirect('logout_message')
        else:
            return redirect(f"{reverse('formacion_opis')}?opi={opi_a_mod}")  # Redirigir a la página anterior

@groups_required('admin', 'formacion')
@login_required
def ver_pdf_opi(request, opi, nombre):
    #Vista para servir el PDF como descarga
    # pdf_path = os.path.join(settings.MEDIA_ROOT, 'firmas_opis', opi, f"{nombre}_acta_formacion_{opi_a_mod}.pdf")
    pdf_path = shared_media_path('firmas_opis', opi, f"{nombre}_acta_formacion_{opi_a_mod}.pdf")
    
    if not os.path.exists(pdf_path):
        raise Http404("El archivo no existe")

    response = FileResponse(open(pdf_path, 'rb'), content_type='application/pdf')
    return response


####################################################################### DEF PARA LA FORMACION COMPLETA #######################################################################################

@groups_required('admin', 'formacion')
@login_required
def formacion_completa(request):
    admin = False
    if request.user.groups.filter(name='admin').exists():
        admin = True
    global puesto_info
    global puestos_dict
    operario_form = OperarioComp(request.GET)
    operarios_info = []
    puesto_form = PuestoComp(request.GET)
    tecnicos_info = []
    puesto_info = []
    formados_info = []
    sin_firma_info = []
    firma_info = []
    puesto_a_buscar = []
    puesto_seleccionado = []
    puesto_traducido = []

    if operario_form.is_valid():
        operario_seleccionado = operario_form.cleaned_data.get('OPERARIO')
        if operario_seleccionado:
            operario_obj = polivalencia.objects.filter(OPERARIO=operario_seleccionado).first()
            if operario_obj:
                campos_excluidos = ['id', 'OPERARIO', 'creado_por', 'creado_en', 'modificado_por', 'modificado_en']
                campos_validos = [
                    (field.name, puestos_dict.get(field.name, field.name))
                    for field in polivalencia._meta.fields
                    if field.name not in campos_excluidos and getattr(operario_obj, field.name) == 1
                ]
                

                operarios_info = {
                    'nombre': operario_obj.OPERARIO,
                    'campos_validos': campos_validos  # lista de tuplas (clave, etiqueta)
                }

    if puesto_form.is_valid():
        puesto_info.clear()
        puesto_seleccionado = puesto_form.cleaned_data.get('PUESTO')

        if puesto_seleccionado:
            puesto_traducido = puestos_dict[puesto_seleccionado]

            # ✅ Filtrar técnicos que SÍ están en formación en el campo correspondiente
            tecnicos_info = polivalencia.objects.filter(
                Q(**{puesto_seleccionado: 1})
            ).values('OPERARIO', puesto_seleccionado)

            # ✅ Crear filtro dinámico SECCION1-SECCION7
            filtro = Q()
            for i in range(1, 8):
                filtro |= Q(**{f'SECCION{i}': puesto_seleccionado})

            opis_filtradas = opis.objects.filter(filtro)

            # ✅ Quedarse solo con versión más alta por base
            import re

            pattern = re.compile(r'^(.*)\[(\d+)\]$')
            latest_versions = {}

            for opi in opis_filtradas:
                match = pattern.match(opi.OPI)
                if match:
                    base = match.group(1)
                    version = int(match.group(2))
                else:
                    base = opi.OPI
                    version = 0

                if base not in latest_versions or version > latest_versions[base][1]:
                    latest_versions[base] = (opi, version)

            # ✅ Guardar solo OPIs más recientes
            puesto_info.extend(v[0].OPI for v in latest_versions.values())

            # ✅ Convertir a string con separador |
            puesto_info = " | ".join(puesto_info)

    
    completa_form = None
    operario_seleccionado = request.GET.get('OPERARIO')
    puesto_seleccionado = request.GET.get('PUESTO')

    completa_form = CompletaForm(operario=operario_seleccionado, puesto=puesto_seleccionado)

    estado_firmas = {
        'firma_alumno': False,
        'firma_formador': False,
        'firma_supervisor': False,
        'firma_dpto': False,
        'PDF': False
    }

    if operario_seleccionado and puesto_seleccionado:
        firma_entry = completa.objects.filter(
            OPERARIO=operario_seleccionado,
            PUESTO=puesto_seleccionado
        ).first()

        if firma_entry and firma_entry.firmas:
            for clave in estado_firmas.keys():
                estado_firmas[clave] = clave in firma_entry.firmas
    
    if estado_firmas['firma_alumno'] and estado_firmas['firma_formador'] and estado_firmas['firma_supervisor'] and estado_firmas['firma_dpto'] and not estado_firmas['PDF']:
        storage = get_messages(request)
        for _ in storage:
            pass
        return render(request, 'generando_pdf.html', {
            'operario': operario_seleccionado,
            'puesto': puesto_seleccionado,
            'opis': puesto_info,
        })

    return render(request, 'form_completa.html', {
        'formados_info': formados_info,
        'sin_firma_info': sin_firma_info,
        'firma_info': firma_info,
        'operario_form': operario_form,
        'operarios_info': operarios_info,
        'puesto_form': puesto_form,
        'tecnicos_info': tecnicos_info,
        'puesto_info': puesto_info,
        'puesto_seleccionado': puesto_seleccionado,
        'puesto_traducido': puesto_traducido,
        'completa_form': completa_form,
        'estado_firmas': estado_firmas,
        'admin': admin,
        'mensaje': messages.get_messages(request)
    })


class PDFConLogo(FPDF):
    def header(self):
        # logo_path = r"C:\sonova\formaciones\media\logo.png"
        logo_path = shared_media_path('logo.png')
        if os.path.exists(logo_path):
            # Puedes ajustar x, y y tamaño a tu gusto
            self.image(logo_path, x=160, y=10, w=35)
        self.set_y(20)


def cargar_teoria(puesto_seleccionado):
    # base_path = r"C:\sonova\formaciones\media\plantillas\teoria"
    base_path = shared_plantillas_path('teoria')
    file_path = os.path.join(base_path, f"{puesto_seleccionado}.txt")

    # Verificar que el archivo existe antes de abrirlo
    if not os.path.exists(file_path):
        return None, "El archivo de preguntas no se encuentra."

    # Leer el contenido del archivo .txt
    with open(file_path, 'r', encoding='utf-8') as file:
        contenido = file.read()

    # Dividir el contenido en bloques por cada pregunta
    bloques_preguntas = contenido.strip().split("\n\n")  # Separar por líneas en blanco

    preguntas = []
    for bloque in bloques_preguntas:
        lineas = bloque.strip().split("\n")  # Dividir cada bloque en líneas
        if len(lineas) < 2:
            continue  # Saltar bloques mal formateados

        # Extraer el número y texto de la pregunta
        match_pregunta = re.match(r'^(\d+)-\s*(.+)', lineas[0])
        if not match_pregunta:
            continue  # Saltar si no coincide el formato esperado
        numero = match_pregunta.group(1)
        texto_pregunta = match_pregunta.group(2)

        # Extraer las opciones
        opciones = []
        respuesta_correcta = None
        for linea in lineas[1:]:
            match_opcion = re.match(r'^([A-H])-\s*(.+?)(\*?)$', linea) #[A-H] para letras de opciones, la H es el máximo
            if match_opcion:
                letra = match_opcion.group(1)
                texto_opcion = match_opcion.group(2).strip()
                es_correcta = match_opcion.group(3) == '*'
                opciones.append({'letra': letra, 'texto': texto_opcion})
                if es_correcta:
                    respuesta_correcta = letra

        preguntas.append({
            'indice': int(numero),
            'texto': texto_pregunta,
            'opciones': opciones,
            'respuesta_correcta': respuesta_correcta
        })

    return preguntas, None

@groups_required('admin', 'formacion', 'tecnicos')
@login_required
def completar_teoria(request):
    global puestos_dict
    if request.method == 'GET':
        operario_nombre = request.GET.get('operario_nombre', '')
        puesto_seleccionado = request.GET.get('puesto_seleccionado', '')
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        puesto_traducido = puestos_dict[puesto_seleccionado]

        # Cargar preguntas
        preguntas, error = cargar_teoria(puesto_seleccionado)
        if error:
            return json_error_detallado(error, 'teoria', puesto_seleccionado)

        contexto = {
            'titulo': f"Examen teórico {puesto_traducido}",
            'operario_nombre': operario_nombre,
            'puesto_seleccionado': puesto_seleccionado,
            'puesto_traducido': puesto_traducido,
            'fecha': fecha_actual,
            'preguntas': preguntas
        }

        return render(request, 'completar_teoria.html', contexto)

    elif request.method == 'POST':
        operario_nombre = request.POST.get('operario_nombre', '')
        puesto_seleccionado = request.POST.get('puesto_seleccionado', '')
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        puesto_traducido = puestos_dict[puesto_seleccionado]

        # Cargar preguntas
        preguntas, error = cargar_teoria(puesto_seleccionado)
        if error:
            return json_error_detallado(error, 'teoria', puesto_seleccionado)

        # Procesar las respuestas del usuario
        respuestas_usuario = {}
        for pregunta in preguntas:
            respuesta = request.POST.get(f"respuesta_{pregunta['indice']}")
            respuestas_usuario[pregunta['indice']] = respuesta

        # Calcular resultados
        correctas = 0
        incorrectas = 0
        for pregunta in preguntas:
            if respuestas_usuario.get(pregunta['indice']) == pregunta['respuesta_correcta']:
                correctas += 1
            else:
                incorrectas += 1

        # Verificar si el examen es apto (más del 80% de aciertos)
        total_preguntas = len(preguntas)
        porcentaje_correctas = (correctas / total_preguntas) * 100

        # Verificar si el examen es apto
        if porcentaje_correctas >= 80:
            teoria = completa.objects.get(OPERARIO=operario_nombre, PUESTO=puesto_seleccionado)
            teoria.TEORIA = True
            teoria.firmas["fecha_teoria"] = fecha_actual
            teoria.firmas["porcentaje_teoria"] = f'{porcentaje_correctas:.0f}%'
            teoria.modificado_por = request.user
            teoria.save()
            # Generar el PDF
            # pdf_path = os.path.join(
            #     r"C:\sonova\formaciones\media\form_completa\teoria_ok",
            #     operario_nombre,
            #     f"{puesto_seleccionado}.pdf"
            # )
            pdf_path = shared_media_path(
                'form_completa',
                'teoria_ok',
                operario_nombre,
                f"{puesto_seleccionado}.pdf"
            )
            os.makedirs(os.path.dirname(pdf_path), exist_ok=True)

            pdf = PDFConLogo()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            pdf.cell(200, 10, txt=f"Evaluación Formación Teoría {puesto_traducido}", ln=True, align='C')
            pdf.cell(200, 10, txt=f"Fecha: {fecha_actual}", ln=True, align='L')
            pdf.cell(200, 10, txt=f"Operario: {operario_nombre}", ln=True, align='L')
            pdf.ln(6)

            for pregunta in preguntas:
                # Mostrar la pregunta
                pdf.multi_cell(0, 10, txt=f"{pregunta['indice']}. {pregunta['texto']}")
                
                # Mostrar las opciones y marcar la seleccionada por el operario
                for opcion in pregunta['opciones']:
                    seleccionada = " <<<" if respuestas_usuario.get(pregunta['indice']) == opcion['letra'] else ""
                    correcta = " (Correcta)" if opcion['letra'] == pregunta['respuesta_correcta'] else ""
                    pdf.cell(0, 10, txt=f"  {opcion['letra']}. {opcion['texto']}{seleccionada}{correcta}", ln=True)
                pdf.ln(1)

            # Añadir una tabla con los resultados al final del PDF
            pdf.ln(10)
            pdf.set_font("Arial", style="B", size=12)
            pdf.cell(0, 10, txt="Resultados del Examen", ln=True, align='C')
            pdf.ln(5)

            # Dibujar la tabla
            pdf.set_font("Arial", size=12)
            column_width = 63  # Ancho de cada columna
            row_height = 10    # Altura de cada fila

            # Primera fila (encabezados)
            pdf.cell(column_width, row_height, "RESPUESTAS CORRECTAS", border=1, align='C')
            pdf.cell(column_width, row_height, "FALLOS", border=1, align='C')
            pdf.cell(column_width, row_height, "RESULTADO", border=1, align='C')
            pdf.ln(row_height)

            # Segunda fila (valores)
            pdf.cell(column_width, row_height, str(correctas), border=1, align='C')
            pdf.cell(column_width, row_height, str(incorrectas), border=1, align='C')
            pdf.cell(column_width, row_height, "PASS", border=1, align='C')
            pdf.ln(row_height)

            pdf.output(pdf_path)

            # Redirigir con mensaje de éxito
            messages.success(request, f"Enhorabuena {operario_nombre}! Has aprobado el examen. Hasta la próxima!")
            #return redirect('formacion_completa')
            base_url = reverse('formacion_completa')  # nombre de la vista
            query_string = urlencode({'OPERARIO': operario_nombre, 'PUESTO': puesto_seleccionado})
            url = f"{base_url}?{query_string}"
            
            if request.user.groups.filter(name='tecnicos').exists():
                return redirect('logout_message')
            else:
                return HttpResponseRedirect(url)
        else:
            # Redirigir con mensaje de fallo
            messages.error(request, f"No apto, lo siento {operario_nombre}, tendrás que volver a intentarlo.")
            #return redirect('formacion_completa')
            base_url = reverse('formacion_completa')  # nombre de la vista
            query_string = urlencode({'OPERARIO': operario_nombre, 'PUESTO': puesto_seleccionado})
            url = f"{base_url}?{query_string}"

            if request.user.groups.filter(name='tecnicos').exists():
                return redirect('logout_message')
            else:
                return HttpResponseRedirect(url)

    # Retorno por defecto si no se cumple ninguna condición
    return HttpResponse("Método no permitido", status=405)

def cargar_practica(puesto_seleccionado):
    # base_path = r"C:\sonova\formaciones\media\plantillas\practica"
    base_path = shared_plantillas_path('practica')
    file_path = os.path.join(base_path, f"{puesto_seleccionado}.txt")

    if not os.path.exists(file_path):
        return None, "El archivo de cuestiones no se encuentra."

    with open(file_path, 'r', encoding='utf-8') as file:
        contenido = file.read()

    bloques_cuestion = contenido.strip().split("\n\n")

    cuestiones = []
    for bloque in bloques_cuestion:
        linea = bloque.strip()
        
        # Buscar formato: número-guion-espacio-texto
        match_cuestion = re.match(r'^(\d+)-\s*(.+)', linea)
        if not match_cuestion:
            continue  # Saltar líneas que no cumplan el formato

        numero = int(match_cuestion.group(1))
        texto = match_cuestion.group(2)

        cuestiones.append({
            'indice': numero,
            'texto': texto,
            'opciones': [
                {'letra': 'A', 'texto': 'Sí'},
                {'letra': 'B', 'texto': 'No'}
            ]
        })

    return cuestiones, None

@groups_required('admin', 'formacion')
@login_required
def completar_practica(request):
    global puestos_dict
    if request.method == 'GET':
        operario_nombre = request.GET.get('operario_nombre', '')
        puesto_seleccionado = request.GET.get('puesto_seleccionado', '')
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        puesto_traducido = puestos_dict[puesto_seleccionado]

        preguntas, error = cargar_practica(puesto_seleccionado)
        if error:
            return json_error_detallado(error, 'practica', puesto_seleccionado)

        contexto = {
            'titulo': f"Evaluación Práctica de {puesto_traducido}",
            'operario_nombre': operario_nombre,
            'puesto_seleccionado': puesto_seleccionado,
            'fecha': fecha_actual,
            'preguntas': preguntas
        }

        return render(request, 'completar_practica.html', contexto)

    elif request.method == 'POST':
        operario_nombre = request.POST.get('operario_nombre', '')
        puesto_seleccionado = request.POST.get('puesto_seleccionado', '')
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        puesto_traducido = puestos_dict[puesto_seleccionado]

        preguntas, error = cargar_practica(puesto_seleccionado)
        if error:
            return json_error_detallado(error, 'practica', puesto_seleccionado)

        respuestas_usuario = {}
        si_contadas = 0

        for pregunta in preguntas:
            respuesta = request.POST.get(f"respuesta_{pregunta['indice']}")
            respuestas_usuario[pregunta['indice']] = respuesta
            if respuesta == 'A':  # 'A' = Sí
                si_contadas += 1

        total_preguntas = len(preguntas)
        porcentaje_si = (si_contadas / total_preguntas) * 100
        no_contadas = total_preguntas - si_contadas  

        if porcentaje_si >= 80:
            practica = completa.objects.get(OPERARIO=operario_nombre, PUESTO=puesto_seleccionado)
            practica.PRACTICA = True
            practica.firmas["fecha_practica"] = fecha_actual
            practica.firmas["porcentaje_practica"] = f'{porcentaje_si:.0f}%'
            practica.modificado_por = request.user
            practica.save()

            # Generar PDF
            pdf_path = shared_media_path(
                'form_completa',
                'practica_ok',
                operario_nombre,
                f"{puesto_seleccionado}.pdf"
            )
            os.makedirs(os.path.dirname(pdf_path), exist_ok=True)

            pdf = PDFConLogo()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            pdf.cell(200, 10, txt=f"Evaluación Formación Práctica {puesto_traducido}", ln=True, align='C')
            pdf.cell(200, 10, txt=f"Fecha: {fecha_actual}", ln=True, align='L')
            pdf.cell(200, 10, txt=f"Operario: {operario_nombre}", ln=True, align='L')
            pdf.ln(6)

            for pregunta in preguntas:
                pdf.multi_cell(0, 10, txt=f"{pregunta['indice']}. {pregunta['texto']}")
                for opcion in pregunta['opciones']:
                    seleccionada = " <<<" if respuestas_usuario.get(pregunta['indice']) == opcion['letra'] else ""
                    pdf.cell(0, 10, txt=f"  {opcion['letra']}. {opcion['texto']}{seleccionada}", ln=True)
                pdf.ln(1)

            pdf.ln(10)
            pdf.set_font("Arial", style="B", size=12)
            pdf.cell(0, 10, txt="Resumen del Cuestionario", ln=True, align='C')
            pdf.ln(5)

            column_width = 63
            row_height = 10
            pdf.set_font("Arial", size=12)
            pdf.cell(column_width, row_height, "SÍ", border=1, align='C')
            pdf.cell(column_width, row_height, "NO", border=1, align='C')
            pdf.cell(column_width, row_height, "RESULTADO", border=1, align='C')
            pdf.ln(row_height)

            pdf.cell(column_width, row_height, str(si_contadas), border=1, align='C')
            pdf.cell(column_width, row_height, str(no_contadas), border=1, align='C')
            pdf.cell(column_width, row_height, "PASS", border=1, align='C')
            pdf.ln(row_height)

            pdf.output(pdf_path)

            messages.success(request, f"¡Buen trabajo {operario_nombre}! Has aprobado la evaluación práctica.")
            #return redirect('formacion_completa')
            base_url = reverse('formacion_completa')  # nombre de la vista
            query_string = urlencode({'OPERARIO': operario_nombre, 'PUESTO': puesto_seleccionado})
            url = f"{base_url}?{query_string}"

            return HttpResponseRedirect(url)
        else:
            messages.error(request, f"No apto, {operario_nombre}. Solo se obtuvieron {porcentaje_si:.0f}% de 'Sí'.")
            #return redirect('formacion_completa')
            base_url = reverse('formacion_completa')  # nombre de la vista
            query_string = urlencode({'OPERARIO': operario_nombre, 'PUESTO': puesto_seleccionado})
            url = f"{base_url}?{query_string}"

            return HttpResponseRedirect(url)

    return HttpResponse("Método no permitido", status=405)

def cargar_producto():
    columnas = [
        'Nº de piezas',
        'Estado pieza general (bueno, regular, malo)',
        'No conformidad (Si/No)',
        'Descripción no conformidad',
        'Tipo no conformidad (Grave, leve)'
    ]

    opciones_estado = ['Bueno', 'Regular', 'Malo']
    opciones_conformidad = ['Sí', 'No']
    opciones_tipo = ['Grave', 'Leve']

    html = '<table border="1" style="border-collapse: collapse; width: 100%;">'
    html += '<thead><tr>'
    for col in columnas:
        html += f'<th>{col}</th>'
    html += '</tr></thead><tbody>'

    for i in range(1, 11):
        html += '<tr>'
        # Nº de piezas
        html += f'<td>{i}</td>'

        # Estado pieza general
        html += '<td><select name="estado_pieza[]">'
        for opcion in opciones_estado:
            html += f'<option value="{opcion.lower()}">{opcion}</option>'
        html += '</select></td>'

        # No conformidad
        html += '<td><select name="no_conformidad[]">'
        for opcion in opciones_conformidad:
            html += f'<option value="{opcion.lower()}">{opcion}</option>'
        html += '</select></td>'

        # Descripción no conformidad
        html += '<td><input type="text" name="descripcion_no_conformidad[]" style="width: 100%;"></td>'

        # Tipo no conformidad
        html += '<td><select name="tipo_no_conformidad[]">'
        for opcion in opciones_tipo:
            html += f'<option value="{opcion.lower()}">{opcion}</option>'
        html += '</select></td>'

        html += '</tr>'

    html += '</tbody></table>'
    return html

@groups_required('admin', 'formacion')
@login_required
def completar_producto(request):
    global puestos_dict
    if request.method == 'GET':
        operario_nombre = request.GET.get('operario_nombre', '')
        puesto_seleccionado = request.GET.get('puesto_seleccionado', '')
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        puesto_traducido = puestos_dict[puesto_seleccionado]

        tabla_html = cargar_producto()

        contexto = {
            'titulo': f"Evaluación Práctica de {puesto_traducido}",
            'operario_nombre': operario_nombre,
            'puesto_seleccionado': puesto_seleccionado,
            'fecha': fecha_actual,
            'tabla_html': tabla_html
        }

        return render(request, 'completar_producto.html', contexto)

    elif request.method == 'POST':
        operario_nombre = request.POST.get('operario_nombre', '')
        puesto_seleccionado = request.POST.get('puesto_seleccionado', '')
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        puesto_traducido = puestos_dict[puesto_seleccionado]

        no_conformidades = request.POST.getlist('no_conformidad[]')
        descripciones = request.POST.getlist('descripcion_no_conformidad[]')
        estados = request.POST.getlist('estado_pieza[]')
        tipos = request.POST.getlist('tipo_no_conformidad[]')

        total_filas = len(no_conformidades)
        cantidad_no = sum(1 for val in no_conformidades if val.lower() == 'no')
        porcentaje_no = (cantidad_no / total_filas) * 100
        

        if porcentaje_no == 100:
            # Aprobado
            producto = completa.objects.get(OPERARIO=operario_nombre, PUESTO=puesto_seleccionado)
            producto.PRODUCTO = True
            producto.firmas["fecha_producto"] = fecha_actual
            producto.firmas["porcentaje_producto"] = f'{porcentaje_no:.0f}%'
            producto.modificado_por = request.user
            producto.save()

            # Generar PDF
            pdf_path = shared_media_path(
                'form_completa',
                'producto_ok',
                operario_nombre,
                f"{puesto_seleccionado}.pdf"
            )
            os.makedirs(os.path.dirname(pdf_path), exist_ok=True)

            pdf = PDFConLogo()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            pdf.cell(200, 10, txt=f"Evaluación Práctica FINAL", ln=True, align='C')
            pdf.cell(200, 10, txt=f"Fecha: {fecha_actual}", ln=True, align='L')
            pdf.cell(200, 10, txt=f"Proceso: {puesto_traducido}", ln=True, align='L')
            pdf.cell(200, 10, txt=f"Operario: {operario_nombre}", ln=True, align='L')
            pdf.ln(10)

            # Tabla de evaluación
            pdf.set_font("Arial", style="B", size=11)
            pdf.cell(20, 10, "Nº", border=1, align='C')
            pdf.cell(40, 10, "Estado", border=1, align='C')
            pdf.cell(35, 10, "No Conf.", border=1, align='C')
            pdf.cell(60, 10, "Descripción", border=1, align='C')
            pdf.cell(35, 10, "Tipo", border=1, align='C')
            pdf.ln()

            pdf.set_font("Arial", size=10)
            for i in range(total_filas):
                pdf.cell(20, 10, str(i+1), border=1, align='C')
                pdf.cell(40, 10, estados[i], border=1, align='C')
                pdf.cell(35, 10, no_conformidades[i], border=1, align='C')

                if no_conformidades[i].lower() == "sí":
                    descripcion = descripciones[i]
                    tipo = tipos[i]
                else:
                    descripcion = ""
                    tipo = ""

                pdf.cell(60, 10, descripcion, border=1, align='C')
                pdf.cell(35, 10, tipo, border=1, align='C')
                pdf.ln()

            pdf.ln(10)
            

            column_width = 63
            row_height = 10
            pdf.set_font("Arial", size=12)
            pdf.cell(column_width, row_height, "PRODUCTOS CORRECTOS", border=1, align='C')
            pdf.cell(column_width, row_height, "NO CONFORMIDADES", border=1, align='C')
            pdf.cell(column_width, row_height, "RESULTADO", border=1, align='C')
            pdf.ln(row_height)

            pdf.cell(column_width, row_height, str(cantidad_no), border=1, align='C')
            pdf.cell(column_width, row_height, str('0'), border=1, align='C')
            pdf.cell(column_width, row_height, "PASS", border=1, align='C')
            pdf.ln(row_height)

            pdf.output(pdf_path)

            messages.success(request, f"¡Buen trabajo {operario_nombre}! Has aprobado la evaluación de producto.")
            #return redirect('formacion_completa')
            base_url = reverse('formacion_completa')  # nombre de la vista
            query_string = urlencode({'OPERARIO': operario_nombre, 'PUESTO': puesto_seleccionado})
            url = f"{base_url}?{query_string}"

            return HttpResponseRedirect(url)

        else:
            messages.error(request, f"No apto, {operario_nombre}. Hay no conformidades detectadas.")
            #return redirect('formacion_completa')
            base_url = reverse('formacion_completa')  # nombre de la vista
            query_string = urlencode({'OPERARIO': operario_nombre, 'PUESTO': puesto_seleccionado})
            url = f"{base_url}?{query_string}"

            return HttpResponseRedirect(url)

    return HttpResponse("Método no permitido", status=405)

@groups_required('admin', 'formacion', 'tecnicos')
@login_required
def firmar_formacion(request):
    if request.method == "GET":
        operario = request.GET.get('operario')
        puesto = request.GET.get('puesto')
        tipo_firma = request.GET.get('tipo_firma')

        return render(request, 'firmar_formacion.html', {
            'operario': operario,
            'puesto': puesto,
            'tipo_firma': tipo_firma
        })
    
@groups_required('admin', 'formacion', 'tecnicos')
@login_required
def guardar_firma(request):
    if request.method == 'POST':
        operario = request.POST.get('operario')
        puesto = request.POST.get('puesto')
        tipo_firma = request.POST.get('tipo_firma')
        imagen_data = request.POST.get('imagen')
        
        if imagen_data:
            try:
                # Ruta de guardado
                # path_firmas = rf"C:\sonova\formaciones\media\form_completa\firmas\{operario}\{puesto}"
                path_firmas = shared_media_path('form_completa', 'firmas', operario, puesto)

                # Crear carpeta si no existe
                os.makedirs(path_firmas, exist_ok=True)

                # Decodificar imagen base64
                format, imgstr = imagen_data.split(';base64,')  # e.g. data:image/png;base64,iVBORw0...
                ext = format.split('/')[-1]  # png

                nombre_archivo = f"{tipo_firma}.png"
                ruta_completa = os.path.join(path_firmas, nombre_archivo)

                with open(ruta_completa, 'wb') as f:
                    f.write(base64.b64decode(imgstr))

                firma = completa.objects.get(OPERARIO=operario, PUESTO=puesto)
                firma.firmas[tipo_firma] = ruta_completa

                if tipo_firma == 'firma_alumno':
                    dni = (request.POST.get('dni') or '').strip().upper()
                    firma.firmas['dni'] = dni
                elif tipo_firma == 'firma_supervisor':
                    supervisor = (request.POST.get('supervisor') or '').strip().upper()
                    firma.firmas['supervisor'] = supervisor
                elif tipo_firma == 'firma_formador':
                    formador = (request.POST.get('formador') or '').strip().upper()
                    firma.firmas['formador'] = formador
                elif tipo_firma == 'firma_dpto':
                    dpto = (request.POST.get('dpto') or '').strip().upper()
                    firma.firmas['dpto'] = dpto

                firma.modificado_por = request.user
                firma.save()

                messages.success(request, "Firma guardada correctamente.")
            except Exception as e:
                messages.error(request, f"Error al guardar la firma: {e}")

        #return redirect('formacion_completa')
        base_url = reverse('formacion_completa')  # nombre de la vista
        query_string = urlencode({'OPERARIO': operario, 'PUESTO': puesto})
        url = f"{base_url}?{query_string}"

        if request.user.groups.filter(name='tecnicos').exists():
            return redirect('logout_message')

    return HttpResponseRedirect(url)  

   
def generar_pdf(request):
    fecha_hora = datetime.now()
    fecha = fecha_hora.strftime("%d-%b-%Y").upper()
    hora = fecha_hora.strftime("%H:%M")
    global puestos_dict
    global puesto_info
    
    excluir_firmas = ['firma_alumno', 'firma_formador', 'firma_supervisor', 'firma_dpto']
    def reemplazar_en_parrafo(paragraph, replacements, excluir_claves=excluir_firmas):
        if excluir_claves is None:
            excluir_claves = []

        texto_original = ''.join(run.text for run in paragraph.runs)
        texto_modificado = texto_original

        for key, value in replacements.items():
            if key not in excluir_claves:
                texto_modificado = texto_modificado.replace(f'{{{{ {key} }}}}', str(value))

        if texto_modificado != texto_original:
            for run in paragraph.runs:
                run.text = ''
            paragraph.runs[0].text = texto_modificado

    operario_nombre = request.POST.get('operario')
    puesto_seleccionado = request.POST.get('puesto')
    puesto_traducido = puestos_dict[puesto_seleccionado]

    # Filtrar por operario y puesto
    filtro = completa.objects.filter(OPERARIO=operario_nombre, PUESTO=puesto_seleccionado).first()

    if not filtro:
        messages.add_message(request, messages.ERROR, "No se encontró la entrada correspondiente.")
        return redirect('formacion_completa')  # Redirigir a algún lugar si no se encuentra

    firmas = filtro.firmas  # El diccionario con las firmas y datos

    # Definir las plantillas .docx para cada tipo de PDF
    plantillas = {
        # 'teoria': r"C:\sonova\formaciones\media\plantillas\validaciones\Validación Initial Training fase 1.docx",
        'teoria': shared_plantillas_path('validaciones', 'Validación Initial Training fase 1.docx'),
        # 'practica': r"C:\sonova\formaciones\media\plantillas\validaciones\Validación Initial Training fase 2.docx",
        'practica': shared_plantillas_path('validaciones', 'Validación Initial Training fase 2.docx'),
        # 'producto': r"C:\sonova\formaciones\media\plantillas\validaciones\Validación Initial Training fase 3.docx",
        'producto': shared_plantillas_path('validaciones', 'Validación Initial Training fase 3.docx'),
        # 'final': r"C:\sonova\formaciones\media\plantillas\validaciones\Validación final.docx",
        'final': shared_plantillas_path('validaciones', 'Validación final.docx'),
        # 'acta': r"C:\sonova\formaciones\media\plantillas\acta_formacion.docx"
        'acta': shared_plantillas_path('acta_formacion.docx')
    }

    # Datos a insertar en cada PDF
    datos = {
        'teoria': {
            'fecha': firmas.get('fecha_teoria'),
            'porcentaje': firmas.get('porcentaje_teoria'),
        },
        'practica': {
            'fecha': firmas.get('fecha_practica'),
            'porcentaje': firmas.get('porcentaje_practica'),
        },
        'producto': {
            'fecha': firmas.get('fecha_producto'),
            'porcentaje': firmas.get('porcentaje_producto'),
        },
        'final': firmas,  # Para el PDF final, usamos todos los datos
        'acta': {
            'opis': puesto_info,
            'dni': firmas.get('dni'),
            'formador': firmas.get('formador'),
            'supervisor': firmas.get('supervisor'),
            'dpto': firmas.get('dpto'),
        },
    }

    # Crear los PDFs
    for tipo, plantilla_path in plantillas.items():
        # Cargar la plantilla .docx
        doc = Document(plantilla_path)

        # Reemplazar los placeholders con los datos correspondientes
        replacements = {
            'fecha': fecha,
            'hora': hora,
            'operario': operario_nombre,
            'dni': firmas.get('dni'),
            'puesto': puesto_traducido,
            'opis': puesto_info,
            'fecha_teoria': firmas.get('fecha_teoria'),
            'porcentaje_teoria': firmas.get('porcentaje_teoria'),
            'fecha_practica': firmas.get('fecha_practica'),
            'porcentaje_practica': firmas.get('porcentaje_practica'),
            'fecha_producto': firmas.get('fecha_producto'),
            'porcentaje_producto': firmas.get('porcentaje_producto'),
            'firma_alumno': firmas.get('firma_alumno'),
            'formador': firmas.get('formador'),
            'firma_formador': firmas.get('firma_formador'),
            'supervisor': firmas.get('supervisor'),
            'firma_supervisor': firmas.get('firma_supervisor'),
            'dpto': firmas.get('dpto'),
            'firma_dpto': firmas.get('firma_dpto'),
        }
        # Agregar los datos específicos (fecha y porcentaje) para cada tipo de PDF
        replacements.update(datos[tipo])
        
        # Reemplazar en los párrafos
        for paragraph in doc.paragraphs:
            reemplazar_en_parrafo(paragraph, replacements, excluir_claves=excluir_firmas)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        reemplazar_en_parrafo(paragraph, replacements, excluir_claves=excluir_firmas)

        def insertar_imagen_en_parrafos(paragraphs, key, image_path):
            if not os.path.isfile(image_path):
                print(f"[ADVERTENCIA] Imagen no encontrada: {image_path}")
                return

            for paragraph in paragraphs:
                full_text = paragraph.text
                if f'{{{{ {key} }}}}' in full_text:
                    # Borrar todo el contenido del párrafo
                    for run in paragraph.runs:
                        run.text = ''
                    # Insertar imagen en un nuevo run
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(1.7))

        for key in ['firma_alumno', 'firma_formador', 'firma_supervisor', 'firma_dpto']:
            firma_path = firmas.get(key)
            if firma_path:
                # Insertar en párrafos normales
                insertar_imagen_en_parrafos(doc.paragraphs, key, firma_path)

                # Insertar en párrafos dentro de tablas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            insertar_imagen_en_parrafos(cell.paragraphs, key, firma_path)


        # Guardar el documento modificado en PDF
        # output_directory = os.path.join(settings.MEDIA_ROOT, 'documentos', operario_nombre, puesto_traducido)
        output_directory = shared_media_path('documentos', operario_nombre, puesto_traducido)
        os.makedirs(output_directory, exist_ok=True)

        output_path = os.path.join(output_directory, f"{operario_nombre}_{tipo}_formacion.docx")
        doc.save(output_path)

        # Llamar a la función para convertir .docx a PDF
        pdf_output_path = output_path.replace(".docx", ".pdf")
        convertir_docx_a_pdf(output_path, pdf_output_path)  # Usamos tu función de conversión

        # Eliminar el archivo .docx después de convertirlo a PDF
        os.remove(output_path)

        # Ruta del PDF combinado
        pdf_final_path = os.path.join(output_directory, f"{puesto_traducido}_formacion_completa.pdf")

        # Inicializa el combinador
        merger = PdfMerger()

        # Orden explícito de los PDFs (tanto generados como preexistentes)
        pdfs_en_orden = [
            os.path.join(output_directory, f"{operario_nombre}_acta_formacion.pdf"),
            os.path.join(output_directory, f"{operario_nombre}_final_formacion.pdf"),
            os.path.join(output_directory, f"{operario_nombre}_teoria_formacion.pdf"),
            # rf"C:\sonova\formaciones\media\form_completa\teoria_ok\{operario_nombre}\{puesto_seleccionado}.pdf",
            shared_media_path('form_completa', 'teoria_ok', operario_nombre, f"{puesto_seleccionado}.pdf"),
            os.path.join(output_directory, f"{operario_nombre}_practica_formacion.pdf"),
            # rf"C:\sonova\formaciones\media\form_completa\practica_ok\{operario_nombre}\{puesto_seleccionado}.pdf",
            shared_media_path('form_completa', 'practica_ok', operario_nombre, f"{puesto_seleccionado}.pdf"),
            os.path.join(output_directory, f"{operario_nombre}_producto_formacion.pdf"),
            # rf"C:\sonova\formaciones\media\form_completa\producto_ok\{operario_nombre}\{puesto_seleccionado}.pdf",            
            shared_media_path('form_completa', 'producto_ok', operario_nombre, f"{puesto_seleccionado}.pdf"),
        ]

        # Añadir los PDFs en orden
        for pdf_path in pdfs_en_orden:
            if os.path.exists(pdf_path):
                merger.append(pdf_path)
            else:
                pass#print(f"[ADVERTENCIA] No se encontró el PDF: {pdf_path}")

        # Guardar el PDF combinado
        with open(pdf_final_path, 'wb') as f_out:
            merger.write(f_out)

        merger.close()

    #borrar archivos innecesarios despues de genear el PDF  
    for path in pdfs_en_orden:
        if os.path.exists(path):
            os.remove(path)

    # shutil.rmtree(rf"C:\sonova\formaciones\media\form_completa\firmas\{operario_nombre}\{puesto_seleccionado}")
    shutil.rmtree(shared_media_path('form_completa', 'firmas', operario_nombre, puesto_seleccionado))
    
    pdf = completa.objects.get(OPERARIO=operario_nombre, PUESTO=puesto_seleccionado)
    pdf.firmas['PDF'] = pdf_final_path
    pdf.creado_por = request.user
    pdf.save()

    tecnico = polivalencia.objects.filter(OPERARIO=operario_nombre).first()
    if tecnico and hasattr(tecnico, puesto_seleccionado):
        if getattr(tecnico, puesto_seleccionado) == 1:
            setattr(tecnico, puesto_seleccionado, 2)
            tecnico.modificado_por = request.user
            tecnico.save()
    
    messages.add_message(request, messages.INFO, "Documento de formación generado correctamente.")
    #return redirect('formacion_completa')  # Redirigir a la página de éxito o donde corresponda
    base_url = reverse('formacion_completa')  # nombre de la vista
    query_string = urlencode({'OPERARIO': operario_nombre, 'PUESTO': puesto_seleccionado})
    url = f"{base_url}?{query_string}"

    return HttpResponseRedirect(url)

@groups_required('admin', 'formacion')
@login_required
def ver_pdf_form(request, nombre, puesto):
    global puestos_dict
    puesto_traducido = puestos_dict[puesto]
    #Vista para servir el PDF como descarga
    # pdf_path = os.path.join(settings.MEDIA_ROOT, 'documentos', nombre, puesto_traducido, f"{puesto_traducido}_formacion_completa.pdf")
    pdf_path = shared_media_path('documentos', nombre, puesto_traducido, f"{puesto_traducido}_formacion_completa.pdf")
    if not os.path.exists(pdf_path):
        raise Http404("El archivo no existe")

    response = FileResponse(open(pdf_path, 'rb'), content_type='application/pdf')
    return response

####################################################################### DEF PARA LA MATRIZ DE POLIVALENCIA #######################################################################################

@groups_required('admin', 'formacion')
@login_required
def editar_matriz(request):
    global puestos_dict
    operario_form = OperarioForm(request.GET)
    operarios_info = []
    puesto_form = PuestoForm(request.GET)
    tecnicos_info = []
    puesto_info = []
    formados_info = []
    nivel = []
    puesto_seleccionado = []
    puesto_traducido = []
    eliminar_form = EliminarForm(request.GET)
    count_formacion = 0
    count_experto = 0
    count_formado = 0
    total_count_formacion = 0
    total_count_experto = 0
    total_count_formado = 0

    if operario_form.is_valid():
        operario_seleccionado = operario_form.cleaned_data.get('OPERARIO')
        if operario_seleccionado:
            operario_obj = polivalencia.objects.filter(OPERARIO=operario_seleccionado).first()
            if operario_obj:
                campos_cero = []
                campos_con_valor = []

                for field in polivalencia._meta.fields:
                    nombre_campo = field.name
                    if nombre_campo in ['id', 'OPERARIO']:
                        continue

                    valor = getattr(operario_obj, nombre_campo)
                    etiqueta = puestos_dict.get(nombre_campo, nombre_campo)

                    if valor == 0:
                        campos_cero.append((nombre_campo, etiqueta))
                    elif valor in [1, 2, 3]:
                        campos_con_valor.append((nombre_campo, etiqueta, valor))

                operarios_info = {
                    'nombre': operario_obj.OPERARIO,
                    'campos_cero': campos_cero,  # lista de tuplas (clave, etiqueta) con valor 0
                    'campos_con_valor': campos_con_valor  # lista de tuplas (clave, etiqueta) con valor 1-3
                }

    if puesto_form.is_valid():
        puesto_info.clear()
        puesto_seleccionado = puesto_form.cleaned_data.get('PUESTO')

        if puesto_seleccionado:
            puesto_traducido = puestos_dict[puesto_seleccionado]

            count_formacion = polivalencia.objects.filter(**{puesto_seleccionado: 1}).count()
            count_experto = polivalencia.objects.filter(**{puesto_seleccionado: 3}).count()
            count_formado = polivalencia.objects.filter(**{puesto_seleccionado: 2}).count()

            # Obtener todos los técnicos con su valor para ese campo
            tecnicos_raw = polivalencia.objects.all().values('OPERARIO', puesto_seleccionado)

            tecnicos_valor_formado = []
            tecnicos_valor_experto = []
            tecnicos_valor_formacion = []


            for tecnico in tecnicos_raw:
                operario = tecnico['OPERARIO']
                valor = tecnico[puesto_seleccionado]

                if valor == 2:
                    tecnicos_valor_formado.append({'OPERARIO': operario, 'valor': valor})
                elif valor == 3:
                    tecnicos_valor_experto.append({'OPERARIO': operario, 'valor': valor})
                elif valor == 1:
                    tecnicos_valor_formacion.append({'OPERARIO': operario, 'valor': valor})

            tecnicos_info = {
                'con_valor_formado': tecnicos_valor_formado,
                'con_valor_experto': tecnicos_valor_experto,
                'con_valor_formacion': tecnicos_valor_formacion

            }
            

    completa_form = None
    operario_seleccionado = request.GET.get('OPERARIO')
    puesto_seleccionado = request.GET.get('PUESTO')

    completa_form = CompletaForm(operario=operario_seleccionado, puesto=puesto_seleccionado)

    if operario_seleccionado and puesto_seleccionado:
        operario_obj = polivalencia.objects.filter(OPERARIO=operario_seleccionado).first()
        if operario_obj:
            nivel = getattr(operario_obj, puesto_seleccionado)

    for campo in puestos_dict.keys():
        total_count_formacion += polivalencia.objects.filter(**{campo: 1}).count()
        total_count_experto += polivalencia.objects.filter(**{campo: 3}).count()
        total_count_formado += polivalencia.objects.filter(**{campo: 2}).count()
    

    return render(request, 'editar_matriz.html', {
        'formados_info': formados_info,
        'operario_form': operario_form,
        'operarios_info': operarios_info,
        'puesto_form': puesto_form,
        'tecnicos_info': tecnicos_info,
        'puesto_info': puesto_info,
        'puesto_seleccionado': puesto_seleccionado,
        'puesto_traducido': puesto_traducido,
        'completa_form': completa_form,
        'nivel': nivel,
        'eliminar_form': eliminar_form,
        'count_formacion': count_formacion,
        'count_experto': count_experto,
        'count_formado': count_formado,
        'total_count_formacion': total_count_formacion,
        'total_count_experto': total_count_experto,
        'total_count_formado': total_count_formado,
        'mensaje': messages.get_messages(request)
    })

@groups_required('admin', 'formacion')
@login_required
def agregar_tecnico(request):
    if request.method == "POST":
        nombre = request.POST.get('nuevo_tecnico', '').upper()
        if nombre:
            # Verificar si ya existe
            if not polivalencia.objects.filter(OPERARIO=nombre).exists():
                nuevo_operario = polivalencia.objects.create(OPERARIO=nombre)
                nuevo_operario.creado_por = request.user
                nuevo_operario.save()
                messages.success(request, f"Técnico '{nombre}' añadido correctamente.")
                return redirect('editar_matriz')
        messages.success(request, f"'{nombre}' ya existe.")        
        return redirect('editar_matriz')

@groups_required('admin', 'formacion')
@login_required    
def eliminar_tecnico(request):
    if request.method == "POST":
        nombre = request.POST.get('ELIMINAR', '').upper()
        if nombre:
            # Verificar si existe
            tecnico = polivalencia.objects.filter(OPERARIO=nombre).first()
            if tecnico:
                tecnico.delete()
                messages.success(request, f"Técnico '{nombre}' eliminado correctamente.")
            else:
                messages.warning(request, f"Técnico '{nombre}' no encontrado.")
        else:        
            messages.warning(request, f"No se ha recibido un nombre válido para eliminar.")
        return redirect('editar_matriz')

@groups_required('admin', 'formacion')
@login_required
def editar_tecnico(request):
    global puestos_dict
    if request.method == "POST":
        operario = request.POST.get('operario', '').upper()
        puesto = request.POST.get('puesto', '')
        puesto_traducido = puestos_dict.get(puesto, puesto)
        nuevo_valor = request.POST.get('nuevo_valor', '')
        if operario and puesto and nuevo_valor.isdigit():
            tecnico = polivalencia.objects.filter(OPERARIO=operario).first()
            if tecnico and hasattr(tecnico, puesto):
                setattr(tecnico, puesto, int(nuevo_valor))
                tecnico.modificado_por = request.user
                tecnico.save()
                if nuevo_valor == '1':
                    nueva_formacion = completa.objects.create(OPERARIO=operario, PUESTO=puesto, PRACTICA=False, PRODUCTO=False, TEORIA=False, firmas={})
                    nueva_formacion.creado_por = request.user
                    nueva_formacion.save()
                messages.success(request, f"Nivel de '{operario}' en '{puesto_traducido}' actualizado correctamente.")
                base_url = reverse('editar_matriz')  # nombre de la vista
                query_string = urlencode({'OPERARIO': operario, 'PUESTO': puesto})
                url = f"{base_url}?{query_string}"

                return HttpResponseRedirect(url)
            else:
                messages.error(request, "No se encontró el técnico o el campo.")
        else:
            messages.error(request, "Datos inválidos para actualizar.")
            base_url = reverse('editar_matriz')  # nombre de la vista
            query_string = urlencode({'OPERARIO': operario, 'PUESTO': puesto})
            url = f"{base_url}?{query_string}"

        return HttpResponseRedirect(url)
    else:
        messages.error(request, "Método no permitido.")
        base_url = reverse('editar_matriz')  # nombre de la vista
        query_string = urlencode({'OPERARIO': operario, 'PUESTO': puesto})
        url = f"{base_url}?{query_string}"

        return HttpResponseRedirect(url)

@groups_required('admin', 'formacion')
@login_required   
def descargar_polivalencia(request):
    # 1) Obtener todos los registros de la base de datos
    registros = polivalencia.objects.all().values()

    # 2) Convertir a DataFrame
    df = pd.DataFrame(registros)

    # 3) Eliminar columna 'id' si existe
    if 'id' in df.columns:
        df = df.drop(columns=['id'])

    # 4) Convertir datetimes con timezone a naive datetimes (sin timezone)
    for col in df.columns:
        if pd.api.types.is_datetime64tz_dtype(df[col]):
            df[col] = df[col].dt.tz_localize(None)

    # 5) Convertir celdas con 0 a vacío y mantener solo 1-3
    def limpiar_valor(x):
        if isinstance(x, (int, float)):
            if x in [1, 2, 3]:
                return x
            else:
                return ""
        return x

    df = df.applymap(limpiar_valor)

    # 6) Generar nombre de archivo con fecha
    fecha_hoy = datetime.now().strftime("%Y-%m-%d")
    filename = f"TPL-708_Matriz_de_Polivalencia_{fecha_hoy}.xlsx"

    # 7) Guardar en memoria
    from io import BytesIO
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Polivalencia')

    # 7) Preparar respuesta de descarga
    output.seek(0)
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

def grafica(request):
    total_count_formacion = 0
    total_count_experto = 0
    total_count_formado = 0

    for campo in puestos_dict.keys():
        total_count_formacion += polivalencia.objects.filter(**{campo: 1}).count()
        total_count_formado += polivalencia.objects.filter(**{campo: 2}).count()
        total_count_experto += polivalencia.objects.filter(**{campo: 3}).count()
        
    return render(request, 'grafica.html' , {
        'total_count_formacion': total_count_formacion,
        'total_count_experto': total_count_experto,
        'total_count_formado': total_count_formado
    })

@groups_required('admin', 'formacion')
@login_required
def auditoria_diaria(request):
    global puestos_dict
    operario_form = OperarioForm(request.GET)
    operarios_info = []
    auditorias_recientes = []
    puesto_form = PuestoForm(request.GET)

    if operario_form.is_valid():
        operario_seleccionado = operario_form.cleaned_data.get('OPERARIO')
        if operario_seleccionado:
            operario_obj = polivalencia.objects.filter(OPERARIO=operario_seleccionado).first()
            if operario_obj:
                campos_excluidos = ['id', 'OPERARIO', 'creado_por', 'creado_en', 'modificado_por', 'modificado_en']
                campos_validos = [
                    (field.name, puestos_dict.get(field.name, field.name))
                    for field in polivalencia._meta.fields
                    if field.name not in campos_excluidos and getattr(operario_obj, field.name) not in [0, 1]
                ]

                operarios_info = {
                    'nombre': operario_obj.OPERARIO,
                    'campos_validos': campos_validos
                }
                
                # Buscar auditorías de los últimos 3 meses
                fecha_limite = timezone.now().date() - timedelta(days=90)
                auditorias_recientes = auditoria.objects.filter(
                    OPERARIO=operario_seleccionado,
                    DIA__gte=fecha_limite
                ).order_by('-DIA')
                
                # Traducir el proceso al nombre legible
                for aud in auditorias_recientes:
                    aud.proceso_traducido = puestos_dict.get(aud.PROCESO, aud.PROCESO)

                return render(request, 'auditoria.html', {
                    'mensaje': messages.get_messages(request),
                    'operario_form': operario_form,
                    'operarios_info': operarios_info,
                    'auditorias_recientes': auditorias_recientes
                })
    
    return render(request, 'auditoria.html', {'mensaje': messages.get_messages(request),
                                                            'operario_form': operario_form,})

@groups_required('admin', 'formacion')
@login_required
def registrar_auditoria(request):
    if request.method == 'POST':
        dia = datetime.now().date()
        sap = request.POST.get('SAP').upper()
        num_serie = request.POST.get('NUM_SERIE').upper()
        familia = request.POST.get('FAMILIA').upper()
        proceso = request.POST.get('PUESTO').upper()
        auditor_nombre = request.POST.get('AUDITOR').upper()
        operario = request.POST.get('OPERARIO').upper()
        no_conformidad = request.POST.get('NO_CONFORMIDAD') == 'on'
        observaciones = request.POST.get('OBSERVACIONES').upper()

        if not no_conformidad:
            observaciones = 'TODO OK'


        if sap and num_serie and familia and proceso and auditor_nombre and operario:
            nueva_auditoria = auditoria.objects.create(
                DIA=dia,
                SAP=sap,
                NUM_SERIE=num_serie,
                FAMILIA=familia,
                PROCESO=proceso,
                AUDITOR=auditor_nombre,
                OPERARIO=operario,
                NO_CONFORMIDAD=no_conformidad,
                OBSERVACIONES=observaciones
            )
            nueva_auditoria.creado_por = request.user
            nueva_auditoria.save()
            messages.success(request, "Auditoría registrada correctamente.")
        else:
            messages.error(request, "Por favor, complete todos los campos obligatorios.")

        

    return redirect('auditoria_diaria')  # Redirigir a la página de auditoría diaria


####################################################################### GESTOR PLANTILLAS .TXT #######################################################################################

PLANTILLAS_ROOT = shared_plantillas_path()


def _safe_path(base, *parts):
    """Devuelve una ruta absoluta dentro de base, o None si hay directory traversal."""
    candidate = os.path.realpath(os.path.join(base, *parts))
    base_real = os.path.realpath(base)
    if candidate.startswith(base_real + os.sep) or candidate == base_real:
        return candidate
    return None


def _listar_directorio(abs_dir, rel_dir):
    """Devuelve listas de subcarpetas y archivos .txt en abs_dir."""
    carpetas = []
    archivos = []
    try:
        for entry in sorted(os.scandir(abs_dir), key=lambda e: (not e.is_dir(), e.name.lower())):
            rel = (rel_dir + '/' + entry.name) if rel_dir else entry.name
            if entry.is_dir():
                carpetas.append({'nombre': entry.name, 'rel': rel})
            elif entry.is_file() and entry.name.lower().endswith('.txt'):
                archivos.append({'nombre': entry.name, 'rel': rel})
    except PermissionError:
        pass
    return carpetas, archivos


def _breadcrumbs(rel_dir):
    """Genera lista de migas de pan [{nombre, rel}] para la ruta relativa actual."""
    if not rel_dir:
        return []
    partes = rel_dir.replace('\\', '/').split('/')
    crumbs = []
    acumulado = ''
    for p in partes:
        acumulado = (acumulado + '/' + p) if acumulado else p
        crumbs.append({'nombre': p, 'rel': acumulado})
    return crumbs


@groups_required('admin', 'formacion')
@login_required
def edicion_plantillas(request):
    base = PLANTILLAS_ROOT

    # Parámetros de navegación
    subdir = request.GET.get('subdir', '').strip().strip('/\\')
    file_rel = request.GET.get('file', '').strip().strip('/\\')

    # Calcular directorio actual
    if subdir:
        abs_dir = _safe_path(base, subdir)
        if not abs_dir or not os.path.isdir(abs_dir):
            messages.error(request, "Carpeta no válida.")
            return redirect('edicion_plantillas')
    else:
        abs_dir = base
        subdir = ''

    # ── Acciones POST ─────────────────────────────────────────────────────────
    if request.method == 'POST':
        accion = request.POST.get('accion')

        # Guardar / crear archivo
        if accion in ('guardar', 'crear'):
            nombre = request.POST.get('nombre_archivo', '').strip()
            contenido = request.POST.get('contenido', '')
            destino_dir_rel = request.POST.get('dir_actual', '')
            if not nombre:
                messages.error(request, "El nombre del archivo no puede estar vacío.")
            else:
                if not nombre.lower().endswith('.txt'):
                    nombre += '.txt'
                abs_file = _safe_path(base, destino_dir_rel, nombre) if destino_dir_rel else _safe_path(base, nombre)
                if not abs_file:
                    messages.error(request, "Ruta no válida.")
                else:
                    try:
                        with open(abs_file, 'w', encoding='utf-8') as f:
                            f.write(contenido)
                        messages.success(request, f"Archivo '{nombre}' guardado correctamente.")
                    except Exception as e:
                        messages.error(request, f"Error al guardar: {e}")
            # Redirigir al directorio actual
            params = f"?subdir={destino_dir_rel}" if destino_dir_rel else ''
            return redirect(f"{reverse('edicion_plantillas')}{params}")

        # Eliminar archivo
        elif accion == 'eliminar':
            file_a_borrar = request.POST.get('file_rel', '').strip()
            abs_file = _safe_path(base, file_a_borrar) if file_a_borrar else None
            dir_despues = request.POST.get('dir_actual', '')
            if not abs_file or not os.path.isfile(abs_file):
                messages.error(request, "Archivo no encontrado.")
            else:
                try:
                    os.remove(abs_file)
                    messages.success(request, f"Archivo '{os.path.basename(abs_file)}' eliminado.")
                except Exception as e:
                    messages.error(request, f"Error al eliminar: {e}")
            params = f"?subdir={dir_despues}" if dir_despues else ''
            return redirect(f"{reverse('edicion_plantillas')}{params}")

    # ── Vista de edición de archivo ───────────────────────────────────────────
    archivo_contexto = None
    if file_rel:
        abs_file = _safe_path(base, file_rel)
        if abs_file and os.path.isfile(abs_file):
            try:
                with open(abs_file, 'r', encoding='utf-8') as f:
                    contenido_archivo = f.read()
                archivo_contexto = {
                    'rel': file_rel,
                    'nombre': os.path.basename(abs_file),
                    'contenido': contenido_archivo,
                    'dir_actual': subdir,
                }
            except Exception as e:
                messages.error(request, f"No se pudo leer el archivo: {e}")
        else:
            messages.error(request, "Archivo no encontrado.")

    # ── Lista de carpetas y archivos ──────────────────────────────────────────
    carpetas, archivos = _listar_directorio(abs_dir, subdir)
    breadcrumbs = _breadcrumbs(subdir)

    return render(request, 'edicion_plantillas.html', {
        'carpetas': carpetas,
        'archivos': archivos,
        'subdir': subdir,
        'breadcrumbs': breadcrumbs,
        'archivo': archivo_contexto,
        'mensaje': messages.get_messages(request),
    })